from django.shortcuts import render, redirect, reverse
from django.http import HttpResponse
from pypdf import PdfReader
from docx import Document
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from colorama import init
import re
from spellchecker import SpellChecker
from .models import ResumeReview

init()

def index(request):
    """
    Render the index page.
    """
    return render(request, 'index.html')

def upload_resume(request):
    """
    Handle resume upload and create ResumeReview object with all analysis data.
    """
    if request.method == 'GET':
        # Render the upload form
        return render(request, 'upload_resume.html')
    elif request.method == 'POST':
        # Handle the file upload
        uploaded_file = request.FILES.get('resume_file')
        job_desc = request.POST.get('job_description', '')

        if uploaded_file:
            # Create ResumeReview object
            resume_review = ResumeReview.objects.create()
            resume_review.resume_file = uploaded_file

            # Process the file based on extension
            ext = uploaded_file.name.lower().split('.')[-1]
            if ext == 'pdf':
                resume_text, has_image, has_table, formatting_issue = process_pdf(resume_review.resume_file)
            elif ext in ['docx', 'doc']:
                resume_text, has_image, has_table, formatting_issue = process_docx(resume_review.resume_file)
            else:
                return HttpResponse("<h1>Unsupported file type.</h1>")

            # Perform all analysis
            spelling_errors = detect_spelling_errors(resume_text)
            missing_sections = detect_missing_sections(resume_text)
            found_sections = [s for s in CRITICAL_SECTIONS if s not in missing_sections]
            date_matches = check_date_formats(resume_text)

            if job_desc:
                keyword_score, matched_keywords = keyword_match(resume_text, job_desc)
            else:
                job_desc = ""
                keyword_score = 100
                matched_keywords = []

            # Generate ATS report
            ats_report = ats_resume_check(
                resume_text, job_desc, has_image, has_table, formatting_issue,
                missing_sections, spelling_errors, date_matches, keyword_score
            )

            # Collect issues found
            issues = []
            if has_image:
                issues.append("Images detected")
            if has_table:
                issues.append("Tables detected")
            if formatting_issue:
                issues.append("More than 3 columns detected")
            if spelling_errors:
                issues.append(f"Spelling errors: {len(spelling_errors)} found")
            if date_matches == 0:
                issues.append("No valid date formats found")

            # Generate recommendations
            recommendations = []
            if has_image:
                recommendations.append("Remove images to improve ATS compatibility")
            if has_table:
                recommendations.append("Convert tables to simple text format")
            if formatting_issue:
                recommendations.append("Simplify formatting to 2-3 columns maximum")
            if missing_sections:
                recommendations.append(f"Add missing sections: {', '.join(missing_sections)}")
            if spelling_errors:
                recommendations.append("Proofread and correct spelling errors")
            if keyword_score < 20:
                recommendations.append("Include more relevant keywords from job description")

            # Create resume stats
            stats = {
                'total_words': len(word_tokenize(resume_text)),
                'spelling_errors': len(spelling_errors),
                'date_matches': date_matches,
                'keyword_matches': len(matched_keywords),
                'has_image': has_image,
                'has_table': has_table,
                'formatting_issues': formatting_issue
            }

            # Fill in all ResumeReview fields
            resume_review.resume_parsed_text = resume_text
            resume_review.overall_score = ats_report['score']
            resume_review.overall_feedback = '; '.join(ats_report['deductions'])
            resume_review.resume_stats = str(stats)
            resume_review.sections_found = ', '.join(found_sections)
            resume_review.missing_sections = ', '.join(missing_sections)
            resume_review.issues_found = '; '.join(issues)
            resume_review.keyword_match_score = ats_report['keyword_score']
            resume_review.recommendations = '; '.join(recommendations)

            # Save the complete review
            resume_review.save()

            # Redirect to review results page
            return redirect(reverse('review_results', args=(resume_review.pk,)))
        else:
            return HttpResponse("<h1>No file uploaded.</h1>")


def review_results(request, pk):
    """
    Render the review results page with saved ResumeReview data.
    """
    try:
        review = ResumeReview.objects.get(pk=pk)
        context = {
            'review': review,
            'sections_found': review.sections_found.split(', ') if review.sections_found else [],
            'missing_sections': review.missing_sections.split(', ') if review.missing_sections else [],
            'issues_found': review.issues_found.split('; ') if review.issues_found else [],
            'recommendations': review.recommendations.split('; ') if review.recommendations else [],
            'stats': eval(review.resume_stats) if review.resume_stats else {}
        }
        return render(request, 'review_results.html', context)
    except ResumeReview.DoesNotExist:
        return HttpResponse("<h1>Review not found.</h1>")


def process_pdf(resume_file):
    """
    Process the uploaded PDF resume file.
    This function should contain the logic to parse the resume and extract relevant information.
    For now, it just returns a placeholder string.
    """
    # Placeholder for actual PDF processing logic
    try:
        reader = PdfReader(resume_file)
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text())
        text = "\n".join(pages)
        has_image, has_table = detect_images_tables_pdf(reader)
        formatting_issue = detect_formatting_issues_pdf(text)
        return text, has_image, has_table, formatting_issue
    except Exception as e:
        return "", False, False, False

def process_docx(resume_file):
    """
    Process the uploaded DOCX resume file.
    This function should contain the logic to parse the resume and extract relevant information.
    For now, it just returns a placeholder string.
    """
    doc = Document(resume_file)
    text_parts = [p.text for p in doc.paragraphs]
    table_text = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                table_text.append(cell.text)
    full_text = '\n'.join(text_parts + table_text)
    has_image, has_table = detect_images_tables_docx(doc)
    formatting_issue = detect_formatting_issues_docx(doc)
    return full_text, has_image, has_table, formatting_issue


CRITICAL_SECTIONS = [
    "education", "experience", "skills", "projects", "contact"
]

MONTH_REGEX = r"(Jan(uary)?|Feb(ruary)?|Mar(ch)?|Apr(il)?|May|Jun(e)?|Jul(y)?|Aug(ust)?|Sep(tember)?|Oct(ober)?|Nov(ember)?|Dec(ember)?)"
YEAR_REGEX = r"\d{4}"

def detect_images_tables_docx(doc: Document):
    has_image = False
    has_table = len(doc.tables) > 0
    # Check for images
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            has_image = True
            break
    return has_image, has_table

def detect_images_tables_pdf(reader: PdfReader):
    # PDF image detection is non-trivial; here, we check for XObject images
    has_image = False
    for page in reader.pages:
        if '/XObject' in page.get('/Resources', {}):
            xobj = page['/Resources']['/XObject']
            for obj in xobj:
                if xobj[obj]['/Subtype'] == '/Image':
                    has_image = True
                    break
    # Table detection in PDF is not reliable without OCR; skip for now
    has_table = False
    return has_image, has_table

def detect_spelling_errors(text):
    spell = SpellChecker()
    words = word_tokenize(text)
    misspelled = spell.unknown([w for w in words if w.isalpha()])
    return list(misspelled)

def detect_formatting_issues_docx(doc: Document):
    # Check for more than 3 columns in any table
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) > 3:
                return True
    return False

def detect_formatting_issues_pdf(text):
    # Heuristic: lines with more than 3 tab-separated columns
    for line in text.splitlines():
        if line.count('\t') > 2:
            return True
    return False

def detect_missing_sections(text):
    found = []
    text_lower = text.lower()
    for section in CRITICAL_SECTIONS:
        if section in text_lower:
            found.append(section)
    missing = [s for s in CRITICAL_SECTIONS if s not in found]
    return missing

def keyword_match(resume_text, job_desc):
    resume_tokens = set([w.lower() for w in word_tokenize(resume_text) if w.isalpha()])
    job_tokens = set([w.lower() for w in word_tokenize(job_desc) if w.isalpha()])
    stop_words = set(stopwords.words('english'))
    resume_tokens = resume_tokens - stop_words
    job_tokens = job_tokens - stop_words
    matched = resume_tokens & job_tokens
    match_score = len(matched) / max(1, len(job_tokens))
    return round(match_score * 100, 2), list(matched)

def check_date_formats(text):
    # Look for Month Year or Month, Year patterns
    pattern = rf"{MONTH_REGEX}[,]?\s+{YEAR_REGEX}"
    matches = re.findall(pattern, text, re.IGNORECASE)
    return len(matches)

def ats_resume_check(resume_text, job_desc, has_image, has_table, formatting_issue, missing_sections, spelling_errors, date_matches, keyword_score):
    # Simple scoring
    score = 100
    deductions = []
    if has_image:
        score -= 10
        deductions.append("Images detected")
    if has_table:
        score -= 10
        deductions.append("Tables detected")
    if formatting_issue:
        score -= 10
        deductions.append("Formatting: More than 3 columns")
    if missing_sections:
        score -= 10 * len(missing_sections)
        deductions.append(f"Missing sections: {', '.join(missing_sections)}")
    if spelling_errors:
        score -= min(10, len(spelling_errors))
        deductions.append(f"Spelling errors: {', '.join(spelling_errors[:5])}")
    if date_matches == 0:
        score -= 5
        deductions.append("No valid date formats found")
    score += int(keyword_score / 10)  # up to +10 for keyword match
    score = max(0, min(100, score))
    return {
        "score": score,
        "deductions": deductions,
        "keyword_score": keyword_score
    }
